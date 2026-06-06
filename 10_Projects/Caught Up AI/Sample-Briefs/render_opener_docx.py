# -*- coding: utf-8 -*-
"""
Caught Up AI - test-batch opener renderer to .docx (for Google Drive / Docs).

One .docx per opener. Each doc is self-contained:
  STUDENT section: clean numbered article, MCQ (no key), discussion, writing prompt.
  TEACHER section: marked article (single YELLOW highlight + inline device label),
    devices-and-why, MCQ + answer key, discussion + sample responses, writing prompt,
    common misconceptions, AP alignment, and a test-run sourcing note.

Conventions carried from render_opener_v2.py and Samuel's 2026-06-06 corrections:
  - ONE highlight color for every device (yellow). One label color (dark blue).
  - Device labels use ONLY the approved AP device vocabulary; name the SPECIFIC device.
  - Highlight the WHOLE device span, not its trigger word.
  - Tag concession alone; do not name refutation unless it is its own standalone sentence.
  - No em-dashes, no emojis.
"""
import os, re, base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
LABEL_RGB = RGBColor(0x1F, 0x4E, 0x79)
GREY_RGB  = RGBColor(0x66, 0x66, 0x66)
MARK = re.compile(r"\[\[(\d+)\]\](.*?)\[\[/\1\]\]", flags=re.S)
DATE = "Saturday, June 6, 2026"

def strip_marks(raw):
    return MARK.sub(lambda m: m.group(2), raw)

def add_body(doc, text, size=11, after=6, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(size)
    r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p

def add_clean_para(doc, raw, idx):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(7)
    n = p.add_run(str(idx) + "  "); n.bold = True; n.font.name = "Times New Roman"; n.font.size = Pt(11)
    r = p.add_run(strip_marks(raw)); r.font.name = "Times New Roman"; r.font.size = Pt(11)
    return p

def add_marked_para(doc, raw, idx, devices):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(7)
    n = p.add_run(str(idx) + "  "); n.bold = True; n.font.name = "Times New Roman"; n.font.size = Pt(11)
    pos = 0
    for m in MARK.finditer(raw):
        if m.start() > pos:
            r = p.add_run(raw[pos:m.start()]); r.font.name = "Times New Roman"; r.font.size = Pt(11)
        k = int(m.group(1)); inner = m.group(2); dev = devices[k-1]["device"]
        hr = p.add_run(inner); hr.font.name = "Times New Roman"; hr.font.size = Pt(11)
        hr.font.highlight_color = WD_COLOR_INDEX.YELLOW
        lab = p.add_run(" [" + dev + "]"); lab.bold = True
        lab.font.name = "Calibri"; lab.font.size = Pt(8); lab.font.color.rgb = LABEL_RGB
        pos = m.end()
    if pos < len(raw):
        r = p.add_run(raw[pos:]); r.font.name = "Times New Roman"; r.font.size = Pt(11)
    return p

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def mcq_block(doc, mcq, with_key=False):
    for i, item in enumerate(mcq, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        r = p.add_run("%d.  %s" % (i, item["stem"])); r.font.name="Times New Roman"; r.font.size=Pt(11)
        for let, o in zip("ABCD", item["options"]):
            op = doc.add_paragraph(); op.paragraph_format.left_indent = Inches(0.3); op.paragraph_format.space_after = Pt(1)
            r = op.add_run("%s)  %s" % (let, o)); r.font.name="Times New Roman"; r.font.size=Pt(11)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    if with_key:
        heading(doc, "Multiple choice answer key", level=2)
        for i, item in enumerate(mcq, 1):
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
            r = p.add_run("%d. Answer %s.  " % (i, item["answer"])); r.bold=True; r.font.name="Times New Roman"; r.font.size=Pt(11)
            r2 = p.add_run(item["rationale"]); r2.font.name="Times New Roman"; r2.font.size=Pt(11)

def build(piece):
    doc = Document()
    base = doc.styles["Normal"]; base.font.name = "Times New Roman"; base.font.size = Pt(11)

    # Masthead
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Caught Up AI"); r.bold=True; r.font.name="Calibri"; r.font.size=Pt(13)
    p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run("The daily AP English Language Opener   .   %s   .   TEST RUN (sample)" % DATE)
    r.font.name="Calibri"; r.font.size=Pt(8.5); r.font.color.rgb = GREY_RGB

    doc.add_heading(piece["headline"], level=0)
    sub = doc.add_paragraph(); sub.paragraph_format.space_after = Pt(8)
    r = sub.add_run("Register: %s" % piece["register"]); r.italic=True; r.font.name="Calibri"; r.font.size=Pt(9); r.font.color.rgb=GREY_RGB

    # ---------- STUDENT ----------
    heading(doc, "Student copy", level=1)
    add_body(doc, "Read the passage, then answer the questions.", size=10, after=8, italic=True, color=GREY_RGB)
    for i, para in enumerate(piece["body"], 1):
        add_clean_para(doc, para, i)
    heading(doc, "Multiple choice", level=2)
    mcq_block(doc, piece["mcq"], with_key=False)
    heading(doc, "Discussion", level=2)
    for i, d in enumerate(piece["discussion"], 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        r = p.add_run("%d.  %s" % (i, d)); r.font.name="Times New Roman"; r.font.size=Pt(11)
    heading(doc, "Writing prompt (%s)" % piece["writing"]["type"], level=2)
    add_body(doc, piece["writing"]["text"])

    doc.add_page_break()

    # ---------- TEACHER ----------
    heading(doc, "Teacher copy", level=1)
    add_body(doc, "Everything in the student copy, plus the marked article, answers, and teaching notes.",
             size=10, after=8, italic=True, color=GREY_RGB)
    heading(doc, "The article, marked for rhetorical devices (one color throughout)", level=2)
    for i, para in enumerate(piece["body"], 1):
        add_marked_para(doc, para, i, piece["devices"])
    heading(doc, "Devices, and why they are here", level=2)
    quotes = {}
    for para in piece["body"]:
        for m in MARK.finditer(para):
            quotes[int(m.group(1))] = m.group(2).strip()
    for i, d in enumerate(piece["devices"], 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5); p.paragraph_format.left_indent = Inches(0.2)
        head = p.add_run(d["device"]); head.bold=True; head.font.color.rgb=LABEL_RGB; head.font.name="Times New Roman"; head.font.size=Pt(11)
        rest = p.add_run(' (paragraph %d): "%s"' % (d["para"], quotes.get(i,""))); rest.font.name="Times New Roman"; rest.font.size=Pt(11)
        p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(7); p2.paragraph_format.left_indent = Inches(0.2)
        r = p2.add_run(d["purpose"]); r.font.name="Times New Roman"; r.font.size=Pt(11)

    heading(doc, "Multiple choice", level=2)
    mcq_block(doc, piece["mcq"], with_key=True)

    heading(doc, "Discussion", level=2)
    for i, d in enumerate(piece["discussion"], 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        r = p.add_run("%d.  %s" % (i, d)); r.font.name="Times New Roman"; r.font.size=Pt(11)
    heading(doc, "Sample strong discussion responses", level=2)
    for i, s in enumerate(piece["sample_responses"], 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
        r = p.add_run("%d.  %s" % (i, s)); r.font.name="Times New Roman"; r.font.size=Pt(11)

    heading(doc, "Writing prompt (%s)" % piece["writing"]["type"], level=2)
    add_body(doc, piece["writing"]["text"])

    heading(doc, "Common misconceptions to watch for", level=2)
    for m in piece["misconceptions"]:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(m); r.font.name="Times New Roman"; r.font.size=Pt(11)

    heading(doc, "AP exam alignment", level=2)
    add_body(doc, piece["ap_alignment"])

    heading(doc, "Test-run sourcing note", level=2)
    add_body(doc, piece["sourcing"], size=10, italic=True, color=GREY_RGB)

    fname = os.path.join(OUT_DIR, "GDoc - %s.docx" % piece["slug"])
    doc.save(fname)
    return fname

# ============================================================ CONTENT

eclipse = {
 "slug":"Totality (Long Look)",
 "headline":"Totality",
 "register":"The Long Look (lyric awe; renders a phenomenon, withholds judgment)",
 "body":[
  "A total solar eclipse is not the sky dimming. [[1]]It is one light traded for another.[[/1]] For an hour the Moon slides across the face of the Sun a little at a time, and almost nothing seems to happen; the day stays bright, because the eye adjusts to the slow theft of light far better than any instrument does, and a Sun nine tenths covered still reads to us as afternoon. Then, in the last minute, the remaining sliver narrows to a thread, [[2]]the thread breaks into a row of separate beads of light where the last sunlight pours through the valleys on the rim of the Moon[[/2]], and the beads wink out one by one.",
  "Then the light that is left is not sunlight. The white face of the Sun is gone behind a black disk, and [[4]]around that disk stands the corona, the Sun's outer atmosphere, a pale and ragged crown of light that no clear day ever lets you see, because the ordinary Sun outshines it a million times over[[/4]]. [[3]]The temperature drops.[[/3]] Birds stop. A false dusk runs around the entire horizon at once, because you are standing in a shadow only about a hundred miles wide and can see daylight past its edge in every direction. Planets come out. The event an astronomer can predict to the second is, for the few minutes it lasts, almost impossible to take in.",
  "None of this would happen without a coincidence of size. The Sun is about four hundred times wider than the Moon and also about four hundred times farther away, so the two disks appear from the ground almost exactly the same size. [[5]]A larger Moon would blot out the corona; a smaller one would leave the Sun's glare and there would be no true totality at all, only a bright ring.[[/5]] The match is not permanent. The Moon drifts farther from the Earth by about an inch and a half each year, and in the deep future the disks will no longer cover. We happen to live in the stretch of time when they do.",
  "Then it ends as fast as it came. A single point of sunlight returns at the far rim, too bright to look at, and the corona disappears into the glare it was always hiding behind. The birds start again. The temperature climbs. Within minutes the day is an ordinary bright afternoon with a bite taken out of the Sun, and the people who drove through the night to stand in the shadow are already arguing about where the next one falls. What they saw was not rare the way a comet is rare; eclipses fall somewhere on Earth every year or two. What is rare is to be standing in the narrow band where the shadow lands, in the few minutes it stays.",
 ],
 "devices":[
  {"para":1,"device":"Metaphor",
   "purpose":"Calling totality a trade of one light for another, rather than a loss of light, frames the whole piece: the eclipse does not subtract the Sun, it swaps the everyday Sun for the hidden corona. The metaphor tells the reader to watch for what appears, not only for what goes dark."},
  {"para":1,"device":"Imagery",
   "purpose":"Exact sensory detail (the thread, the separate beads, sunlight pouring through valleys on the Moon's rim) lets the reader see the phenomenon known as Baily's beads without being told to feel anything. In this register the wonder is carried by the accuracy of the picture, not by adjectives."},
  {"para":2,"device":"Telegraphic sentence",
   "purpose":"Three words, set hard against the long sentences around them, enact the suddenness they name. The clipped line drops the temperature as fast as the sentence ends and launches the run of short sentences that follow."},
  {"para":2,"device":"Cumulative and periodic sentences",
   "purpose":"The sentence keeps adding members (the corona, what it is, what it looks like, why we never see it), so its growing length mirrors the thing unfolding overhead. The final clause delivers the fact that earns the awe, that the Sun normally outshines the corona a million times over, with no emotional word attached."},
  {"para":3,"device":"Antithesis",
   "purpose":"Two opposed hypotheticals, a larger Moon set against a smaller one, are balanced to show that totality rides on a knife-edge of size. The contrast makes a dry astronomical coincidence feel precarious, and therefore remarkable."},
 ],
 "mcq":[
  {"stem":"In the third paragraph, the writer explains the matched apparent sizes of the Sun and Moon primarily to",
   "options":[
    "account for why totality, and the corona it reveals, are possible at all",
    "argue that the eclipse is evidence of deliberate cosmic design",
    "warn the reader that total eclipses will soon stop occurring",
    "compare the eclipse with other rare astronomical events"],
   "answer":"A",
   "rationale":"The paragraph gives the cause of totality: because the disks appear the same size, the Moon can cover the bright face yet leave the corona showing. B imports a claim of design the piece never makes (it withholds judgment). C names a real detail (the Moon's recession) but the paragraph's purpose is explanation, not warning. D states a true effect tied to the wrong span: the comet comparison is in the fourth paragraph, not the third."},
  {"stem":"The writer's use of very short sentences such as \"The temperature drops\" and \"Birds stop\" serves primarily to",
   "options":[
    "register the abrupt, physical suddenness of totality against the long sentences around them",
    "suggest that the writer cannot find words adequate to the experience",
    "summarize the scientific causes of the temperature change",
    "build suspense about whether totality will occur"],
   "answer":"A",
   "rationale":"The clipped sentences fall right as totality begins, so their speed enacts the sudden change and contrasts with the surrounding long sentences. B reads the brevity too literally as a confession of inadequacy. C names a true topic but the wrong function: the short lines state effects, they do not explain causes. D inverts the moment: totality is already underway, so there is no suspense left to build."},
 ],
 "discussion":[
  "The piece never states what an eclipse means or how you should feel about it, yet it clearly wants you to feel something. How does the writer create awe while withholding any direct statement of significance?",
  "Find a place where the writer admits a limit, something hard to take in or impossible to see normally. What does that candor add to the writer's authority as an observer?",
 ],
 "sample_responses":[
  "A strong answer notes that the awe is built entirely from accurate description and rhythm: the beads, the corona, the false dusk, the short sentences snapping against long ones. Because the writer never says the eclipse is amazing, the reader supplies the feeling, which lands harder than an instruction to feel it would. Withholding the moral is the technique, not an omission.",
  "The writer admits the corona is something \"no clear day ever lets you see\" and that the predictable event is \"almost impossible to take in.\" Conceding the limits of normal sight and of comprehension makes the observer sound honest rather than promotional, so the wonder reads as earned by the facts rather than sold.",
 ],
 "writing":{"type":"in-class, Q2 rhetorical analysis",
  "text":"In a well-developed essay, analyze the rhetorical choices the writer makes to convey the experience of a total solar eclipse. Refer to specific choices in the passage (for example sentence length, imagery, and contrast) and explain how they work."},
 "misconceptions":[
  "Students may call \"It is one light traded for another\" personification. Nothing is given human traits; it is a metaphor, an implied identity between the eclipse and a trade.",
  "Students may treat the third paragraph as the writer's argument or thesis. This register renders rather than argues; there is no claim being defended, only a cause being explained, so do not ask students to find a position the writer takes.",
  "Students may label the beads description a simile. There is no \"like\" or \"as\"; it is imagery, direct sensory description.",
 ],
 "ap_alignment":"Q2 rhetorical-analysis value: a compact text for teaching high-burst rhythm (short flat sentences set against long cumulative ones), precise imagery, and antithesis, inside a register that produces awe without ever asserting that something is significant. It is a clean contrast piece to set beside an argumentative opener, since students must analyze effect without a thesis to lean on.",
 "sourcing":"Test-run sourcing note: the astronomy here is long-established and stable, not pegged to a breaking event. For publication each fact would still pass the standard check: the roughly four-hundred-times coincidence of size and distance; Baily's beads; the corona and its brightness relative to the Sun's face (about a million to one); the umbral shadow roughly a hundred miles wide; the Moon's recession of about an inch and a half a year; eclipses occurring somewhere on Earth every year or two.",
}

books = {
 "slug":"The Books We Don't Read (Long Think)",
 "headline":"The Books We Don't Read",
 "register":"The Long Think (a mind reasoning to an earned turn; no biography)",
 "body":[
  "Almost everyone who buys books owns more of them than they will ever read. [[1]]The unread ones gather on the shelf in a faint atmosphere of reproach, and most of us have made our peace with a low and manageable guilt about them: the history begun and abandoned, the novel everyone discussed that we never opened, the serious book bought in a serious mood and shelved in a lighter one.[[/1]] The guilt assumes that a book does its job only while it is being read, and that an unread book is a small standing failure.",
  "There is something to that. [[2]]A book you keep meaning to read can be a quiet lie you tell about yourself, a picture of the person you would be if you had the evenings.[[/2]] Shelves fill with intentions that were never quite serious. If you have never finished a single book you own, then owning a thousand of them proves nothing except that you can buy things.",
  "[[3]]But the guilt mistakes what the shelf is for.[[/3]] The writer Nassim Nicholas Taleb gave the unread shelf a name, the antilibrary, after the enormous unread collection of Umberto Eco, and made a small durable case for it: the books you have already read are the least useful part of a library. What you have read has done its work and become, at best, a part of you, and at worst a thing you half remember. It is the unread spines that keep the more valuable knowledge in front of you, which is the knowledge of how much there is that you do not know.",
  "Read that way, the unread shelf turns from an accusation into an instrument. [[4]]A shelf of finished books is a record of where you have been; a shelf of unfinished ones is the part that keeps you honest.[[/4]] It is [[5]]a map of a country you have not visited and now at least know is there[[/5]]. The aim is not to read more, though you might. The aim is to keep the size of your own ignorance somewhere you can see it.",
  "None of this makes the guilt entirely wrong, only early. A few of the books really are pure pretension, and you will probably never read those either. But the rest are not waiting to shame you. They are holding your place in a conversation much larger than you are, the one you wandered into when you started buying books you were not ready for, and have not finished wandering through.",
 ],
 "devices":[
  {"para":1,"device":"Cumulative and periodic sentences",
   "purpose":"The sentence accumulates examples under one frame (the history, the novel, the serious book), so the shape of the prose performs the way unread books themselves pile up. The colon turns a general claim into a concrete catalogue the reader recognizes from a real shelf."},
  {"para":2,"device":"Concession",
   "purpose":"Before turning, the writer grants the guilt its real force: some unread books really are small vanities. Conceding this in good faith is what earns the right to overturn the guilt a paragraph later. Without the concession, the reversal would read as easy reassurance."},
  {"para":3,"device":"Shift / volta",
   "purpose":"The turn is this whole clause, not the word \"But\" that opens it. It pivots the essay from agreeing with the reader's guilt to redefining what the shelf is for, and every sentence after it argues from this new footing."},
  {"para":4,"device":"Parallelism",
   "purpose":"Two clauses in the same grammatical shape set the finished shelf against the unfinished one, so the reader weighs them as a matched pair. The balance makes the surprising claim, that the unfinished shelf is the honest one, land as if it were obvious."},
  {"para":4,"device":"Metaphor",
   "purpose":"Recasting the unread shelf as a map of an unvisited country turns a private guilt into a usable picture: the unread books mark the edges of what you know. Held back until the turn is complete, the metaphor is what makes the argument stick."},
 ],
 "mcq":[
  {"stem":"In the second paragraph, which of the following best characterizes the writer's position on the guilt readers feel about unread books?",
   "options":[
    "It is partly warranted, because some unread books reflect intentions that were never serious",
    "It is entirely unwarranted and ought to be dismissed at once",
    "It proves that buying books is a waste of money",
    "It explains why people tend to read less as they grow older"],
   "answer":"A",
   "rationale":"The paragraph grants the guilt real ground (\"There is something to that\") while limiting it, which is a partial endorsement. B overstates by making the guilt entirely baseless, the opposite of \"there is something to that.\" C reads one line too literally and inflates it into a sweeping claim the writer does not make. D imports an outside explanation the passage never raises."},
  {"stem":"The relationship between the second and third paragraphs is best described as one in which the writer first",
   "options":[
    "grants the guilt real force and then argues that it misreads the shelf's purpose",
    "defines a key term and then illustrates that term with an example",
    "recalls a personal memory and then draws a general lesson from it",
    "raises an objection and then concedes that the objection is correct"],
   "answer":"A",
   "rationale":"Paragraph two concedes the guilt; paragraph three turns with \"But the guilt mistakes what the shelf is for\" and reframes the shelf. B mislabels the structure: paragraph two is a concession, not a definition. C is the tempting trap, but this register carries no autobiography, so there is no personal memory. D inverts the move, reversing concession and turn."},
 ],
 "discussion":[
  "The writer spends a paragraph agreeing that the guilt over unread books is partly deserved before turning against it. How does granting that ground change the way the later claim lands?",
  "The essay redefines a shelf of unread books as \"a map of a country you have not visited.\" What does that metaphor let you see that the plain word \"unread\" does not?",
 ],
 "sample_responses":[
  "A strong answer sees the concession as strategic. By admitting that some unread books are vanities, the writer removes the easy objection that this is just comfort for people who buy too much. Having been fair to the guilt, the later claim that the unread shelf is the useful one reads as an earned reversal rather than a rationalization.",
  "Calling the shelf a map turns a feeling (guilt over what is unread) into a tool (a picture of what you do not yet know). \"Unread\" measures only failure against an intention; \"a map of a country you have not visited\" makes the same books point outward, toward the size of one's ignorance, which is the writer's actual point.",
 ],
 "writing":{"type":"homework or extended in-class, Q3 argument",
  "text":"The writer claims that what you have not yet read can be more valuable to keep in view than what you have. Defend, challenge, or qualify the claim that an awareness of our own ignorance is worth preserving. Use your own knowledge, observation, or reading."},
 "misconceptions":[
  "Students may read \"a map of a country you have not visited\" as a simile. There is no \"like\" or \"as\"; it is a metaphor.",
  "Students may take the concession in the second paragraph as the writer's settled view. It is granted on purpose, in order to turn; the writer's real position arrives only after the volta.",
  "Students may call \"But the guilt mistakes what the shelf is for\" a simple transition. It is the volta, the hinge of the whole piece; the connective \"But\" only signals the turn, while the turn itself is the claim that follows.",
  "Students may hunt for the writer's personal story. This register reasons without autobiography, so there is no \"I remember\" scene to find; the evidence is the reasoning, not a life.",
 ],
 "ap_alignment":"Q3 argument and Q2 analysis value: a clean model of concession-then-turn, a single load-bearing metaphor held back until the turn, the parallelism that frames the central reversal, and the withheld-turn architecture. Useful for teaching students to locate a volta as a claim rather than a signal word, and to notice an argument that reasons rather than confesses.",
 "sourcing":"Test-run sourcing note: the one external attribution is the \"antilibrary,\" a term from Nassim Nicholas Taleb's The Black Swan (2007), describing the large unread collection of Umberto Eco. Confirm the attribution and wording before any publication. The remainder is reflection on a durable theme, not reported current-events fact.",
}

recycling = {
 "slug":"The Recycling We Were Sold (Reckoning)",
 "headline":"The Recycling We Were Sold",
 "register":"The Reckoning (moral charge to a public; nameable party, datable trigger)",
 "body":[
  "Start with the number, because the number is the whole case. About nine percent of all the plastic the world has ever thrown away has actually been recycled, according to the OECD; the rest, more than nine tenths, was buried, burned, or left to crumble into pieces too small to collect. For more than thirty years we were sold a different story, told not in figures but in a symbol: an arrow bent into a triangle, stamped on the bottom of a yogurt cup, promising that if you rinsed it and sorted it the rest would follow.",
  "It did not follow, and the people who printed that arrow knew it might not. In 2020 a joint investigation by NPR and the PBS program Frontline reported that plastic makers had promoted recycling for decades while doubting, in their own internal words, that most plastic could ever be recycled at a profit. Look at what the arrangement asks of each side. [[1]]A few dozen companies make the plastic, sell it, and keep the profit; a few hundred million of us rinse it, sort it, carry it to the curb, and feel for a moment that we have done our part.[[/1]] [[2]]The work was handed down; the responsibility was handed off.[[/2]]",
  "And the cost does not wait for us to notice it. While we are sorting, the plastic that was never going to be recycled is already going somewhere: [[3]]when a full truck leaves the curb and is buried instead of remade, when a bale is shipped across an ocean to a country with no way to process it and burned in the open air, when a bag breaks down not into nothing but into fragments that turn up in tap water, in rain, and in human blood[[/3]], the promise on the bottom of the cup is worse than unmet. It has become cover for the very thing it claimed to prevent.",
  "This is not a failure of householders. People did the small thing they were asked to do, faithfully, for a generation. The failure belongs to the companies that built a symbol to look like a solution, and to the public officials who let the symbol stand in for a system. So keep sorting your recycling; some of it is real, and the habit is not the problem.",
  "The problem is who gets to decide how big the problem is. So ask for the number, not the symbol. Ask the company that sells you the cup what becomes of it after the curb, and make an honest answer the price of your business. The arrow was printed to turn a public failure into a private chore. [[4]]Take it back.[[/4]]",
 ],
 "devices":[
  {"para":2,"device":"Antithesis",
   "purpose":"The sentence sets a few dozen producers against a few hundred million sorters in matched form, so the reader feels the imbalance the whole piece is about: profit on one side, unpaid labor and false comfort on the other. The opposition is the engine of the argument, not an ornament."},
  {"para":2,"device":"Parallelism",
   "purpose":"Two short clauses in the same shape, handed down against handed off, compress the entire accusation into one line. The matched grammar makes the swap audible: the same gesture moved the work one way and the responsibility the other."},
  {"para":3,"device":"Anaphora",
   "purpose":"Three clauses beginning with \"when\" pile concrete instances of harm under one repeated frame, building pressure toward the main clause that follows them. The repetition sits at the START of each clause, which is what makes this anaphora and not parallelism, and each \"when\" raises the scale, from one truck to an ocean to the human body."},
  {"para":5,"device":"Telegraphic sentence",
   "purpose":"Three words deliver the charge the whole piece has loaded. After the long accumulating sentences, the abruptness is the point: the reader is handed one plain instruction, and the piece stops before it can soften."},
 ],
 "mcq":[
  {"stem":"Which of the following best describes the writer's primary purpose in the opener?",
   "options":[
    "To move responsibility for plastic waste from households back to the companies that marketed recycling",
    "To persuade readers that they should stop sorting their recycling entirely",
    "To explain the chemical process by which plastics break down over time",
    "To mourn the decline of a recycling system that once worked well"],
   "answer":"A",
   "rationale":"The piece absolves householders (\"This is not a failure of householders\") and assigns the failure to producers and officials, which is a reassignment of responsibility. B inverts a stated point: the writer says \"keep sorting your recycling.\" C names a real detail (plastic crumbling) but not the purpose. D imports a premise the piece denies, since it argues the system never worked as promised."},
  {"stem":"In the third paragraph, the repeated \"when\" clauses serve primarily to",
   "options":[
    "accumulate concrete instances of harm that build pressure toward the writer's charge",
    "list the steps a householder should follow in order to recycle correctly",
    "concede that recycling does sometimes work as it was intended to",
    "supply the statistical evidence behind the nine percent figure"],
   "answer":"A",
   "rationale":"The \"when\" clauses stack escalating harms (a buried truckload, an ocean shipment burned abroad, fragments in the body) to drive toward the indictment. B reads them as instructions, which they are not. C inverts their force; they show failure, not success. D ties them to the wrong span: the statistic appears in the first paragraph, and these clauses are images of consequence, not data."},
 ],
 "discussion":[
  "The writer insists \"This is not a failure of householders\" before assigning blame elsewhere. Why might the writer absolve the reader first, and how does that choice affect the force of the charge that follows?",
  "The opener leads with a single statistic, \"about nine percent,\" before any argument. How does putting the number first shape the way you read everything after it?",
 ],
 "sample_responses":[
  "A strong answer sees the absolution as strategic. By telling readers their sorting was honest work, the writer keeps them from feeling attacked and frees them to direct anger at the producers instead. The charge at the end (\"Take it back\") then lands as shared cause rather than personal scolding, which is exactly the stance the register wants.",
  "Leading with \"about nine percent\" sets a fact as the floor of the argument before any emotion enters, so the reader judges the rest against a number rather than a mood. It also creates the gap the piece exploits: a tiny real figure against a thirty-year promise, so everything after the statistic reads as the distance between what we were told and what was true.",
 ],
 "writing":{"type":"homework or extended in-class, Q3 argument",
  "text":"The writer argues that the people who create a problem should not be the ones who decide how large that problem is. Defend, challenge, or qualify this claim, using your own knowledge, observation, or reading."},
 "misconceptions":[
  "Students may say the writer wants people to stop recycling. The writer explicitly says \"keep sorting your recycling\"; the target is the producers and the symbol, not the householder's habit.",
  "Students may treat \"Take it back\" as merely the ending. In this register it is the charge the whole piece is built to deliver, the imperative the runway and the statistic have been loading.",
  "Students may call the few-against-many contrast a simple comparison. It is antithesis: two opposed sides set in balanced form, and it carries the argument rather than decorating it.",
  "Students may label the \"when ... when ... when\" run parallelism. Because the repeated word opens each clause, it is anaphora; reserve parallelism for repeated structure that does not begin the clause, like \"handed down ... handed off.\"",
 ],
 "ap_alignment":"Q3 argument value: a hot persuasive register for teaching the periodic anaphoric runway answered by a short button, antithesis as the engine of an argument, a statistic placed first as moral fact, and a charge-to-act close. Set beside a descriptive or analytic opener, it shows students how register changes the writer's relationship to the reader, from observer to advocate.",
 "sourcing":"Test-run sourcing note: three claims must clear the Accuracy-Guardrail before this could publish. (1) \"About nine percent\" of plastic waste recycled: OECD Global Plastics Outlook (2022); confirm the figure and its exact scope (annual plastic waste versus all plastic ever made, which is a separate Geyer et al. 2017 finding). (2) The 2020 NPR and PBS Frontline investigation reporting that the plastics industry promoted recycling while internally doubting it: confirm the outlets, the date, and the internal-doubt characterization against the original reporting. (3) Microplastics found in tap water, rain, and human blood: confirm against primary studies (for human blood, the 2022 Environment International study). Treat all three as flagged until independently re-sourced.",
}

PIECES = [eclipse, books, recycling]

if __name__ == "__main__":
    for p in PIECES:
        f = build(p)
        b = base64.b64encode(open(f, "rb").read()).decode()
        with open(f + ".b64", "w") as fh:
            fh.write(b)
        print("wrote:", os.path.basename(f), "| bytes", os.path.getsize(f), "| b64", len(b))
